import datetime

from PyQt6 import QtWidgets
from PyQt6 import QtCore
from PyQt6.QtWidgets import QApplication, QCheckBox, QFileDialog, QInputDialog, QLabel, QLineEdit, QMainWindow, QPushButton
from PyQt6.QtWidgets import QVBoxLayout, QWidget, QHBoxLayout, QPlainTextEdit, QMessageBox, QMenuBar
from PyQt6.QtWebEngineWidgets import QWebEngineView
import language_tool_python
from highlighter_markdown import MarkdownHighlighter
from highlighter_python import PythonHighlighter
from highlighter_html import HtmlHighlighter
from markdown_renderer import render_markdown
from e2 import CodeEditor
import re, yaml
from outline_panel import OutlinePanel 
from pathlib import Path
from PyQt6.QtCore import QRegularExpression, QUrl, Qt
from PyQt6.QtGui import QColor, QFont, QImage, QGuiApplication, QAction, QTextCharFormat, QTextCursor
from PyQt6.QtCore import QStandardPaths
import os
from frontmatter_panel import FrontmatterPanel
import pdfplumber
import pytesseract
from PIL import Image, ImageDraw, ImageFont
from langdetect import detect, LangDetectException
#import language_tool_python
from PyQt6.QtGui import QTextDocument
from PyQt6.QtPrintSupport import QPrinter

from docx import Document
from bs4 import BeautifulSoup
from markdown import markdown

import hashlib
from ebooklib import epub

from markdownify import markdownify as md

from _fontsize import fontsize_counts

# instellingen importeren
from config import FRONTMATTER_TEXT, configuratie, font_sizes
from memo import Memo
from memolijst import MemoLijst



class Markdown_Editor(QWidget):
    def __init__(self):
        super().__init__()

        self.unsaved_changes = False
        self.current_path = None

        LANG_MAP = {
            "nl": "nld",
            "en": "eng",
            "fr": "fra",
            "de": "deu",
            "es": "spa",
            "it": "ita",
        }
        

        # menu begin
        actie = {}

        def maak_menu_punt(self, naam_actie, naam_in_menu, sneltoets, functie):
            actie[naam_actie] = QAction(naam_in_menu, self)
            if len(sneltoets) > 1:
                actie[naam_actie].setShortcut(sneltoets)
            actie[naam_actie].triggered.connect(functie)
            return

        # Maak acties        
         
        # Bestand - Nieuw, Openen, Opslaan, Opslaan als, Sluiten

        maak_menu_punt(self, "nieuw_actie", "Nieuw", "Ctrl+N", self.nieuw)
        
        maak_menu_punt(self, "openen_actie", "Openen...", "Ctrl+O", self.openen)

        maak_menu_punt(self, "invoegen_actie", "Invoegen...", "", self.invoegen)

        maak_menu_punt(self, "opslaan_actie", "Opslaan", "Ctrl+S", self.opslaan)

        maak_menu_punt(self, "opslaan_als_actie", "Opslaan als...", "Ctrl+Alt+S", self.opslaan_als)

        maak_menu_punt(self, "importeer_pdf_als_tekst_actie", "Importeer pdf als tekst", "", self.import_pdf_as_text)

        maak_menu_punt(self, "importeer_pdf_als_md_actie", "Importeer pdf als markdown", "", self.import_pdf_as_md)

        maak_menu_punt(self, "importeer_epub_actie", "Importeer ePub", "", self.import_epub)

        maak_menu_punt(self, "export_pdf_actie", "Exporteer als PDF", "", self.export_pdf)

        maak_menu_punt(self, "export_word_actie", "Exporteer naar Word", "", self.export_word)

        maak_menu_punt(self, "export_epub_actie", "Exporteer naar ePub", "", self.export_epub)

        maak_menu_punt(self, "export_txt_actie", "Exporteer als tekst", "", self.export_txt)

        maak_menu_punt(self, "afsluiten_actie", "Afsluiten", "Ctrl+Q", self.afsluiten)
 
        # Bewerken - Kopieren, Plakken, Knippen, Zoeken, Alles selecteren, Ongedaan maken, Opnieuw doen,
        #     Normaliseren, Geen hoofdletters, Schrift

        maak_menu_punt(self, "kopieren_actie", "Kopieren", "Ctrl+C", self.kopieren)

        maak_menu_punt(self, "plakken_actie", "Plakken", "Ctrl+V", self.plakken)

        maak_menu_punt(self, "knippen_actie", "Knippen", "Ctrl+X", self.knippen)

        maak_menu_punt(self, "zoeken_actie", "Zoeken...", "Ctrl+F", self.zoeken)

        maak_menu_punt(self, "zoeken_en_vervangen_actie", "Zoeken en vervangen...", "Ctrl+H", self.zoeken_en_vervangen)

        maak_menu_punt(self, "alles_selecteren_actie", "Alles selecteren", "Ctrl+A", self.alles_selecteren)

        maak_menu_punt(self, "ongedaan_maken_actie", "Ongedaan maken", "Ctrl+Z", self.ongedaan_maken)

        maak_menu_punt(self, "opnieuw_doen_actie", "Opnieuw doen", "Ctrl+R", self.opnieuw_doen)

        maak_menu_punt(self, "normaliseren_actie", "Normaliseren", "Alt+N", self.normaliseren)

        maak_menu_punt(self, "geen_hoofdletters_actie", "Geen hoofdletters", "Alt+U", self.geen_hoofdletters)

        maak_menu_punt(self, "schrift_actie", "Schrift", "Alt+S", self.schrift)

        maak_menu_punt(self, "spellcheck_actie", "Spelling controleren", "F7", self.spellcheck)

        # Beeld - Lichte modus, Donkere modus, Blauwe modus, Font, Lettergrootte

        maak_menu_punt(self, "lichte_modus_actie", "Lichte modus", "", self.lichte_modus)

        maak_menu_punt(self, "donkere_modus_actie", "Donkere modus", "", self.donkere_modus)

        maak_menu_punt(self, "blauwe_modus_actie", "Blauwe modus", "", self.blauwe_modus)

        maak_menu_punt(self, "font_actie", "Font...", "", self.font)

        maak_menu_punt(self, "favoriete_font_actie", "Favoriete font", "", self.favoriete_font)

        # Navigatie

        maak_menu_punt(self, "naar_begin_actie", "Naar begin", "Ctrl+Home", self.naar_begin)

        maak_menu_punt(self, "naar_einde_actie", "Naar einde", "Ctrl+End", self.naar_einde)

        # Invoegen - Datum, Tijd, md link, md afbeelding, if name == main, frontmatter

        maak_menu_punt(self, "datum_actie", "Datum", "Alt+D", self.datum)

        maak_menu_punt(self, "tijd_actie", "Tijd", "Alt+T", self.tijd)

        maak_menu_punt(self, "md_link_actie", "md link", "Alt+L", self.md_link)

        maak_menu_punt(self, "md_afbeelding_actie", "md afbeelding", "Alt+A", self.md_afbeelding)

        maak_menu_punt(self, "if_name_is_main_actie", "if name == main", "Alt+I", self.if_name_is_main)

        maak_menu_punt(self, "frontmatter_actie", "Frontmatter", "Alt+F", self.frontmatter)

        # Apps - Memo, Memolijst

        maak_menu_punt(self, "memo_actie", "Memo", "", self.memo)

        maak_menu_punt(self, "memolijst_actie", "Memolijst", "", self.memolijst)

        # Help - Over Edith, Sneltoetsen, Sneltoetsen (Alt), Markdown
        
        maak_menu_punt(self, "over_actie", "Over Edith", "", self.over)

        maak_menu_punt(self, "sneltoetsen_actie", "Sneltoetsen", "", self.sneltoetsen)

        maak_menu_punt(self, "sneltoetsen_alt_actie", "Sneltoetsen (Alt)", "", self.sneltoetsen_alt)

        maak_menu_punt(self, "markdown_actie", "Markdown", "", self.markdown_overzicht)
        
        # Maak menubalk en menu's        
        
        menubalk = QMenuBar(self)                 # maak menubalk widget
        menubalk.setStyleSheet("padding: 8px;")
        menubalk.setMinimumHeight(30)
        menubalk.setMaximumHeight(60)
        
        bestand_menu = menubalk.addMenu("Bestand")        
        bestand_menu.addAction(actie["nieuw_actie"])        
        bestand_menu.addAction(actie["openen_actie"])
        bestand_menu.addAction(actie["invoegen_actie"])
        bestand_menu.addAction(actie["opslaan_actie"]) 
        bestand_menu.addAction(actie["opslaan_als_actie"])
        bestand_menu.addSeparator()
        bestand_menu.addAction(actie["importeer_pdf_als_tekst_actie"])
        bestand_menu.addAction(actie["importeer_pdf_als_md_actie"])
        bestand_menu.addAction(actie["importeer_epub_actie"])
        bestand_menu.addSeparator()
        bestand_menu.addAction(actie["export_pdf_actie"])
        bestand_menu.addAction(actie["export_word_actie"])
        bestand_menu.addAction(actie["export_epub_actie"])
        bestand_menu.addAction(actie["export_txt_actie"])
        bestand_menu.addSeparator()
        bestand_menu.addAction(actie["afsluiten_actie"])

        bewerken_menu = menubalk.addMenu("Bewerken")
        bewerken_menu.addAction(actie["kopieren_actie"])
        bewerken_menu.addAction(actie["knippen_actie"])
        bewerken_menu.addAction(actie["plakken_actie"])
        bewerken_menu.addSeparator()
        # bewerken_menu.addAction(actie["zoeken_actie"])
        # bewerken_menu.addAction(actie["zoeken_en_vervangen_actie"])
        # bewerken_menu.addSeparator()
        bewerken_menu.addAction(actie["alles_selecteren_actie"])
        bewerken_menu.addAction(actie["ongedaan_maken_actie"])
        bewerken_menu.addAction(actie["opnieuw_doen_actie"])
        bewerken_menu.addSeparator()
        bewerken_menu.addAction(actie["normaliseren_actie"])
        bewerken_menu.addAction(actie["geen_hoofdletters_actie"])
        bewerken_menu.addAction(actie["schrift_actie"])
        bewerken_menu.addSeparator()
        bewerken_menu.addAction(actie["spellcheck_actie"])

        beeld_menu = menubalk.addMenu("Beeld")
        beeld_menu.addAction(actie["lichte_modus_actie"])
        beeld_menu.addAction(actie["donkere_modus_actie"])
        beeld_menu.addAction(actie["blauwe_modus_actie"])
        beeld_menu.addSeparator()
        beeld_menu.addAction(actie["font_actie"])
        beeld_menu.addAction(actie["favoriete_font_actie"])
        
        navigatie_menu = menubalk.addMenu("Navigatie")
        navigatie_menu.addAction(actie["naar_begin_actie"])
        navigatie_menu.addAction(actie["naar_einde_actie"])

        invoegen_menu = menubalk.addMenu("Invoegen")
        invoegen_menu.addAction(actie["datum_actie"])
        invoegen_menu.addAction(actie["tijd_actie"])
        invoegen_menu.addAction(actie["md_link_actie"])
        invoegen_menu.addAction(actie["md_afbeelding_actie"])
        invoegen_menu.addAction(actie["if_name_is_main_actie"])
        invoegen_menu.addAction(actie["frontmatter_actie"])

        apps_menu = menubalk.addMenu("Apps")
        apps_menu.addAction(actie["memo_actie"])
        apps_menu.addAction(actie["memolijst_actie"])
        
        hulp_menu = menubalk.addMenu("Help")        
        hulp_menu.addAction(actie["over_actie"])
        hulp_menu.addAction(actie["sneltoetsen_actie"])
        hulp_menu.addAction(actie["sneltoetsen_alt_actie"])
        hulp_menu.addAction(actie["markdown_actie"])

        v_layout = QVBoxLayout(self)
        v_layout.setMenuBar(menubalk)  

        self.file_label = QLabel("?")  # bestandsnaam
        self.file_label.setStyleSheet("padding: 8px;")
        self.file_label.setMinimumHeight(30)
        self.file_label.setMaximumHeight(60)


        v_layout.addWidget(self.file_label, 0)


        # menu einde

        layout = QHBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)


        v_layout.addLayout(layout)
        #self.frontmatter = FrontmatterPanel(self)
        #layout.addWidget(self.frontmatter, 0)

        self.statusbar = QtWidgets.QStatusBar()
        self.statusbar.setStyleSheet("padding: 8px;")
        self.statusbar.setMinimumHeight(30)
        self.statusbar.setMaximumHeight(60)

        # zoken en vervangen panel
        # Zoek- en vervang widgets
        find_label = QLabel("Zoeken:")
        self.find_input = QLineEdit()
        self.case_cb = QCheckBox("Hoofdlettergevoelig")
        next_btn = QPushButton("Volgende")
        prev_btn = QPushButton("Vorige")

        replace_label = QLabel("Vervangen door:")
        self.replace_input = QLineEdit()
        replace_btn = QPushButton("Vervangen")
        replace_all_btn = QPushButton("Vervang alles")

        # Layout
        top_layout = QHBoxLayout()
        top_layout.addWidget(find_label)
        top_layout.addWidget(self.find_input)
        top_layout.addWidget(self.case_cb)
        top_layout.addWidget(prev_btn)
        top_layout.addWidget(next_btn)

        bottom_layout = QHBoxLayout()
        bottom_layout.addWidget(replace_label)
        bottom_layout.addWidget(self.replace_input)
        bottom_layout.addWidget(replace_btn)
        bottom_layout.addWidget(replace_all_btn)

        main_layout = QVBoxLayout()
        #main_layout.addWidget(self.editor)
        main_layout.addLayout(top_layout)
        main_layout.addLayout(bottom_layout)

        v_layout.addLayout(main_layout)

        # Signalen
        next_btn.clicked.connect(self.find_next)
        prev_btn.clicked.connect(self.find_previous)
        replace_btn.clicked.connect(self.replace_one)
        replace_all_btn.clicked.connect(self.replace_all)
        self.find_input.textChanged.connect(self.update_highlight)
        self.case_cb.stateChanged.connect(self.update_highlight)

        # Highlight format
        self.highlight_format = QTextCharFormat()
        self.highlight_format.setBackground(QColor("#FFF59D"))  # lichtgeel
        self.match_format = QTextCharFormat()
        self.match_format.setBackground(QColor("#FFCC80"))  # geselecteerde match iets donkerder
        #

        v_layout.addWidget(self.statusbar, 0)
        self.statusbar.showMessage("Ready", 3000)

        # Editor links
        self.editor = CodeEditor()
        self.editor.set_highlighter(MarkdownHighlighter)
        layout.addWidget(self.editor, 1)

        self.editor.textChanged.connect(self.on_text_changed)

        # Outline panel
        self.outline_panel = OutlinePanel()
        layout.addWidget(self.outline_panel, 0)
        self.outline_panel.itemClicked.connect(self.jump_to_heading)  # _panel


        # Preview rechts
        self.preview = QWebEngineView()
        layout.addWidget(self.preview, 1)

        # Live updates
        self.editor.textChanged.connect(self.update_preview)

        # Initial render
        self.update_preview()

        self.editor.verticalScrollBar().valueChanged.connect(self.sync_scroll_to_preview)


    def on_text_changed(self):
        self.unsaved_changes = True
        self.statusbar.showMessage("Onopgeslagen wijzigingen")

    def _editor_scroll_ratio(self):
        sb = self.editor.verticalScrollBar()
        if sb.maximum() == 0:
            return 0
        return sb.value() / sb.maximum()
    
    def sync_scroll_to_preview(self):
        ratio = self._editor_scroll_ratio()
        
        js = f"""
        (function() {{
            let h = document.body.scrollHeight - window.innerHeight;
            window.scrollTo(0, h * {ratio});
        }})();
        """

        self.preview.page().runJavaScript(js)

    def update_preview(self):
        text = self.editor.toPlainText()
        html = render_markdown(text)
        ratio = self._editor_scroll_ratio()

        #self.frontmatter.load_frontmatter(text)

        # bepaal directory van het huidige bestand
        base_path = getattr(self, "current_file_path", None)
        if base_path:
            base_url = QUrl((base_path))  # assets map in dezelfde directory)
        else:
            base_url = QUrl("file:///")  # fallback

        self.preview.setHtml(html, baseUrl=base_url)

        # Sroll herstellen zodra de pagina geladen is
        def after_load(_):
            """
            Dit berekent de maximale verticale scrollafstand (in pixels) die de gebruiker kan scrollen op de pagina.

            Uitleg kort:
            - document.body.scrollHeight = totale hoogte van de inhoud van de pagina.
            - window.innerHeight = hoogte van het zichtbare venster (viewport).
            - h = document.body.scrollHeight - window.innerHeight = 
            het verschil tussen inhoudshoogte en viewport‑hoogte → de maximale scrollTop-waarde (0 tot h). 
            Als h ≤ 0 is, is er niks om te scrollen.
            """
            js = f"""
            (function() {{
                let h = document.body.scrollHeight - window.innerHeight;
                window.scrollTo(0, h * {ratio});
            }})();
            """
            self.preview.page().runJavaScript(js)

        self.preview.page().loadFinished.connect(lambda _: after_load(None))

        self.outline_panel.update_outline(self.parse_headings())
        
    def parse_headings(self):
        text = self.editor.toPlainText().split("\n")
        headings = []

        for i, line in enumerate(text):
            match = re.match(r'^(#{1,6})\s+(.*)', line)
            if match:
                level = len(match.group(1))
                title = match.group(2).strip()
                headings.append((level, title, i))

        return headings
    
    def jump_to_heading(self, item):
        line = item.data(32)  # Retrieve line number from user data
        cursor = self.editor.textCursor()
        cursor.movePosition(cursor.MoveOperation.Start)
        cursor.movePosition(cursor.MoveOperation.Down, n=line)
        self.editor.setTextCursor(cursor)
        self.editor.setFocus()

    def insertFromMimeData(self, source):
        # Controleer of de geplakte data een afbeelding bevat
        if source.hasImage():
            self._handle_paste_image(source)
            return
        else:
            super().insertFromMimeData(source)

    def _handle_paste_image(self, source):
        image = source.imageData()
        if not isinstance(image, QImage):
            return

        # Bepaal waar het huidige markdown-bestand staat
        editor = self.parent().parent() # Assuming the parent of Markdown_Editor is the main window
        base_path = getattr(editor, "current_file_path", None)

        if base_path:
            # geen bestand geopend, gebruik standaard afbeeldingenmap
            save_dir = Path(QStandardPaths.writableLocation(QStandardPaths.StandardLocation.PicturesLocation))
            
        else:
            # Gebruik dezelfde map als het markdown-bestand
            save_dir = Path(base_path).parent / "assets"

        save_dir.mkdir(parents=True, exist_ok=True)

        # Genereer een unieke bestandsnaam
        filename = f"pasted_{QGuiApplication.applicationPid()}_{id(image)}.png"
        file_path = save_dir / filename

        # Sla de afbeelding op
        image.save(str(file_path))

        # Voeg markdown-syntax toe voor de afbeelding
        rel_path = file_path.relative_to(Path(base_path).parent) if base_path else file_path
        self.insertPlainText(f"![{filename}]({rel_path.as_posix()})")

    def load_markdown_file(self, file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
            self.editor.setPlainText(content)
            self.current_file_path = file_path
            self.update_preview()

    def to_file_uri(path_or_uri):    
        if path_or_uri.startswith("file://"):        
            return path_or_uri    
        return Path(path_or_uri).resolve().as_uri()
    
    def update_frontmatter_from_panel(self, new_data):
        text = self.editor.toPlainText()

        # Vervang bestaande frontmatter
        import yaml
        new_yaml = yaml.safe_dump(new_data, sort_keys=False).strip()

        new_frontmatter = f"---\n{new_yaml}\n---"
        
        # Vervang in document
        import re
        updated = re.sub(r"^---\n(.*?)\n---", new_frontmatter, text, flags=re.DOTALL)
        
        self.editor.blockSignals(True)
        self.editor.setPlainText(updated)
        self.editor.blockSignals(False)

        # Preview opnieuw renderen
        self.update_preview()

    def load_highlighter(self):
        path = ""
        i = self.current_path.rfind(".")  # index van de laatste punt       
        if i != -1:
            path = self.current_path[i:]

        if path == ".py":
            self.editor.set_highlighter(PythonHighlighter)
        elif path == ".md":
            self.editor.set_highlighter(MarkdownHighlighter)
        elif path == ".html" or path == ".htm":
            self.editor.set_highlighter(HtmlHighlighter)
        else:
            self.editor.set_highlighter(MarkdownHighlighter)

    # menu acties

    def nieuw(self):        
        if self.unsaved_changes:
            reply = QMessageBox.question(self, 'Waarschuwing', 
                                         'Huidig bestand is nog niet opgeslagen. Wil je de wijzigingen opslaan?')
            if reply == QMessageBox.StandardButton.Yes:
                self.opslaan()
        self.editor.clear()
        self.setWindowTitle("Geen naam")
        self.current_path = None 
        self.statusbar.showMessage("Nieuw Bestand")
        self.file_label.setText("?")

    def openen(self):        
        if self.unsaved_changes:
            reply = QMessageBox.question(self, 'Waarschuwing', 
                                         'Huidig bestand is nog niet opgeslagen. Wil je de wijzigingen opslaan?')
            if reply == QMessageBox.StandardButton.Yes:
                self.opslaan()
        try:
            fname = QFileDialog.getOpenFileName(self, 'Open bestand', configuratie["opslaglocatie"], 'Alle bestanden (*)')
            self.setWindowTitle(fname[0])
            self.file_label.setText(fname[0])
            
            with open(fname[0], 'r') as f:
                filetext = f.read()
                self.editor.setPlainText(filetext)
            self.current_path = fname[0]
            self.load_highlighter()
            self.statusbar.showMessage("Bestand geopend")
        except Exception as e:
            self.dialog_critical(str(e))

    def invoegen(self):
        try:
            fname = QFileDialog.getOpenFileName(self, 'Tekst of markdown bestand invoegen', configuratie["opslaglocatie"], 'Alle bestanden (*)')
            with open(fname[0], 'r') as f:
                filetext = f.read()
                self.editor.insertPlainText(filetext)
            self.statusbar.showMessage("Bestand ingevoegd")
        except Exception as e:
            self.dialog_critical(str(e))

    def opslaan(self):        
        if self.current_path is not None:
            filetext = self.editor.toPlainText()
            try:
                with open(self.current_path, 'w') as f:
                    f.write(filetext)
                self.statusbar.showMessage("Bestand opgeslagen")
                self.file_label.setText(self.current_path)
            except Exception as e:
                self.dialog_critical(str(e))
        else:
            self.opslaan_als()

    def opslaan_als(self):        
        try:
            pathname = QFileDialog.getSaveFileName(self, 'Bestand opslaan', configuratie["opslaglocatie"], 'Tekst bestanden (*.txt)')
            filetext = self.editor.toPlainText()
            with open(pathname[0], 'w') as f:
                f.write(filetext)
            self.current_path = pathname[0]
            self.setWindowTitle(pathname[0])
            self.file_label.setText(pathname[0])
            self.statusbar.showMessage("Bestand opgeslagen")
        except Exception as e:
            errortekst = "Bestand niet opgeslagen!\n" + str(e)
            self.dialog_critical(errortekst)

    def afsluiten(self):
        QMessageBox.information(self, "Afsluiten", "Programma afsluiten.")
        # waarschuwen bij onopgeslagen wijzigingen
        if self.unsaved_changes:
            reply = QMessageBox.question(self, 'Waarschuwing', 
                                         'Huidig bestand is nog niet opgeslagen. Wil je de wijzigingen opslaan?')
            if reply == QMessageBox.StandardButton.Yes:
                self.opslaan()
        CodeEditor.close(self)
        # Alle vensters sluiten
        for widget in QApplication.topLevelWidgets():
            widget.close()

    def kopieren(self):
        self.editor.copy()

    def knippen(self):
        self.editor.cut()

    def plakken(self):
        self.editor.paste()

    def zoeken(self):
        text, ok = QInputDialog.getText(self, 'Zoeken', 'Voer de zoekterm in:')
        if ok and text:
            gevonden = self.editor.find(text)
            if not gevonden:
                # weer naar boven
                cursor = self.editor.textCursor()
                cursor.movePosition(QTextCursor.MoveOperation.Start)
                self.editor.setTextCursor(cursor)
                gevonden = self.editor.find(text)
                if not gevonden:
                    term = text.replace('<','&lt;')
                    term = term.replace('>','&gt;')
                    QMessageBox.information(self, "Vinden", f"'{term}' niet gevonden")

    def alles_selecteren(self):
        self.editor.selectAll()

    def ongedaan_maken(self):
        self.editor.undo()

    def opnieuw_doen(self):
        self.editor.redo()

    def normaliseren(self):
        cursor = self.editor.textCursor()
        volledige_tekst = self.editor.toPlainText()
        tussenstap = ' '.join(volledige_tekst.split('\n'))
        genormaliseerde_tekst = '.\n'.join(tussenstap.split('.'))
        self.editor.setPlainText(genormaliseerde_tekst)

    def geen_hoofdletters(self):
        selectie = self.editor.textCursor()
        if selectie.hasSelection():
            geselecteerde_tekst = selectie.selectedText()
            kleine_tekst = geselecteerde_tekst.lower()
            selectie.insertText(kleine_tekst)
        else:
            QMessageBox.about(self, "Geen Selectie", "Selecteer eerst tekst om om te zetten naar kleine letters.")

    def schrift(self):
        selectie = self.editor.textCursor()
        if selectie.hasSelection():
            geselecteerde_tekst = selectie.selectedText()
            schrift_tekst = ''
            for char in geselecteerde_tekst:
                if 'A' <= char <= 'Z':
                    schrift_char = chr(ord(char) + 0x1D4D0 - ord('A'))
                elif 'a' <= char <= 'z':
                    schrift_char = chr(ord(char) + 0x1D4EA - ord('a'))
                else:
                    schrift_char = char
                schrift_tekst += schrift_char
            selectie.insertText(schrift_tekst)
        else:
            QMessageBox.about(self, "Geen Selectie", "Selecteer eerst tekst om om te zetten naar schrift.")

    def lichte_modus(self):
        configuratie["darkmode"] = 'light'
        self.setStyleSheet('')

    def donkere_modus(self):
        configuratie["darkmode"] = 'dark'
        self.setStyleSheet('''
            QWidget{
                background-color: rgb(33,33,33);
                color: #FFFFFF;
                }
            QPlainTextEdit{
                background-color: rgb(46,46,46);
                color: #FFFFFF;
            }
            ''')

    def blauwe_modus(self):
        configuratie["darkmode"] = 'blue'
        self.setStyleSheet('''
                background-color: #0000AA;
                color: #FFFFFF;
                QWidget{
                    background-color: #0000AA;
                    color: #FFFFFF;
                    }
                QPlainTextEdit{  
                    background-color: #000BFF;
                    color: #FFFFFF;
                    }
                QMenuBar{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QMenu{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QStatusBar{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QLabel{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QPushButton{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QLineEdit{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QCheckBox{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QScrollBar{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QHeaderView::section{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QTreeView{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QTableView{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QTabWidget::pane{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QTabBar::tab{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QTabBar::tab:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QScrollBar::handle{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QScrollBar::add-line, QScrollBar::sub-line{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QScrollBar::add-page, QScrollBar::sub-page{
                    background-color: #0000AA;
                    color: #FFFFFF;
                }
                QMenuBar::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QMenu::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QStatusBar::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QLabel::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QPushButton::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QLineEdit::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QCheckBox::item:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                QScrollBar::handle:selected{
                    background-color: #000BFF;
                    color: #FFFFFF;
                }
                ''')

    def font(self):
        from PyQt6.QtWidgets import QFontDialog
        font, ok = QFontDialog.getFont()
        if ok:
            self.editor.setFont(font)
            print(f"Font ingesteld op: {font.family()}, grootte: {font.pointSize()}")
            print(font.toString())

    def set_font(self, font_name, font_size):
        font = QFont(font_name, font_size)
        self.editor.setFont(font)

    def favoriete_font(self):
        font_name, font_size = configuratie['favoriete_font']
        self.set_font(font_name, font_size)

    def naar_begin(self):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        self.editor.setTextCursor(cursor)

    def naar_einde(self):
        cursor = self.editor.textCursor()
        cursor.movePosition(QTextCursor.MoveOperation.End)
        self.editor.setTextCursor(cursor)

    def over(self):        
        QMessageBox.information(self, "Over Edith", "Markdown editor met preview.")

    def sneltoetsen(self):
        QMessageBox.about(self, "Sneltoetsen", "Ctrl+N: Nieuw\nCtrl+O: Openen\nCtrl+S: Opslaan\n" \
        "Ctrl+Shift+S: Opslaan als\nCtrl+Q: Sluiten\nCtrl+C: Kopiëren\nCtrl+X: Knippen\nCtrl+V: Plakken\n" \
        "Ctrl+F: Zoeken\nCtrl+Z: Ongedaan maken\nCtrl+R: Opnieuw doen\nCtrl+A: Alles selecteren")

    def sneltoetsen_alt(self):
        QMessageBox.about(self, "Sneltoetsen Alt", "Invoegen:\nAlt+D: Datum\nAlt+T: Tijd\nAlt+L: md link\n" \
                          "Alt+A: md afbeelding\nAlt+I: if name == main\nAlt+F: Frontmatter\n\n" \
                          "Bewerken:\nAlt+N: Normaliseren\nAlt+U: Geen hoofdletters\nAlt+S: Schrift\n\n")


    def markdown_overzicht(self):
        QMessageBox.about(self, "Markdown", 
                          "Koppen:\n" \
                          "# H1\n## H2\n### H3\n\n" \
                          "Vet (bold):\t**vet**\n\n" \
                          "Schuin (italic):\t*schuin*\n\n" \
                          "Blockquote:\n> blockquote\n\n" \
                          "Genummerde lijst:\n1. eerste item\n2. tweede item\n3. derde item\n\n" \
                          "Ongenummerde lijst:\n- item A\n- item B\n- item C\n\n" \
                          "Highlight:\t==highlighted==\n")

    def datum(self):
        nu = datetime.datetime.now()
        datum_nu = nu.strftime("%Y-%m-%d")
        self.editor.insertPlainText(datum_nu)

    def tijd(self):
        nu = datetime.datetime.now()
        tijd_nu = nu.strftime("%H:%M")
        self.editor.insertPlainText(tijd_nu)

    def md_link(self):
        pathname = QFileDialog.getOpenFileName(self, 'Bestand openen', configuratie["opslaglocatie"], 'Alle bestanden (*)')
        if pathname[0]:
            bestandsnaam = pathname[0].split('/')[-1]
            md_code = f"[{bestandsnaam}]({pathname[0]})"
            self.editor.insertPlainText(md_code)

    def md_afbeelding(self):
        pathname = QFileDialog.getOpenFileName(self, 'Afbeelding openen', configuratie["opslaglocatie"], 'Afbeeldingen (*.png *.jpg *.jpeg *.bmp *.gif);;Alle bestanden (*)')
        if pathname[0]:
            bestandsnaam = pathname[0].split('/')[-1]
            md_code = f"![{bestandsnaam}]({pathname[0]})"
            self.editor.insertPlainText(md_code)

    def if_name_is_main(self):
        self.editor.insertPlainText("if__name__ == '__main__':\n    ")

    def frontmatter(self):
        fm = FRONTMATTER_TEXT
        self.editor.insertPlainText(fm)


    def memo(self):
        self.memo_venster = Memo()
        self.memo_venster.show()

    def memolijst(self):
        self.memo_lijst_venster = MemoLijst()
        self.memo_lijst_venster.show()

    def dialog_critical(self, s):
        dlg = QMessageBox(self)
        dlg.setText(s)
        dlg.setIcon(QMessageBox.Icon.Critical)
        dlg.show()

    # self.import_pdf_as_text
    # > self.pdf_to_text(path)

    # self.import_pdf_as_md


    def pdf_to_text(self, path):
        text = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text.append(page.extract_text() or "")
        return "\n\n".join(text)
    
    def open_pdf_as_text(self, path):
        text = self.pdf_to_text(path)
        self.editor.setPlainText(text)

    def open_pdf_as_markdown(self, path):
        text = self.pdf_to_text(path)
        md = text.replace("\n", "  \n")
        self.editor.setPlainText(md)

    def import_pdf_as_text(self):
        path, _ = QFileDialog.getOpenFileName(self, "Kies een pdf om te importeren", "", "PDF-bestanden (*.pdf)")
        if not path:
            return
        try:
            text = self.pdf_to_text(path)  # self.pdf_to_text_with_ocr(path)
        except Exception as e:
            QMessageBox.critical(self, "Fout bij importeren", str(e))
            return
        # Plaats tekst in editor
        self.editor.setPlainText(text)

        self.current_file_path = None
        self.file_label.setText("?")

    def import_pdf_as_md(self):
        path, _ = QFileDialog.getOpenFileName(self, "Kies een pdf om te importeren", "", "PDF-bestanden (*.pdf)")
        if not path:
            return
        try:
            text = self.pdf_to_markdown(path)  # self.pdf_to_markdown_with_ocr
        except Exception as e:
            QMessageBox.critical(self, "Fout bij importeren", str(e))
            return
        # Plaats tekst in editor
        self.editor.setPlainText(text)

        self.current_file_path = None
        self.file_label.setText("?")

    def pdf_to_markdown(self, path: str) -> str:
        # fontsizes
        # counts_by_page, dict(total_counter) = fontsize_counts(path)
        per_page, totaal  = fontsize_counts(path)

        # Sorteer op fontsize oplopend
        grootste_font_aantal = 0
        grootste_font_size = 0
        for size in sorted(totaal):
            fontmaat = size
            fontaantal = totaal[size]
            
            if fontaantal > grootste_font_aantal:
                grootste_font_aantal = fontaantal
                grootste_font_size = fontmaat


        kop_fonts = {}
        for size in sorted(totaal):
            if size > grootste_font_size:
                kop_fonts[size] = totaal[size]
        lengte = len(kop_fonts)
        volgende = "H1"
        for fontmaat in (sorted(kop_fonts, reverse=True)):
            if lengte > 0:
                font_sizes[volgende] = fontmaat - 0.1
                if volgende == "H3":
                    lengte = 0
                if volgende == "H2":
                    volgende = "H3"
                if volgende == "H1":
                    volgende = "H2"





        #
        md_lines = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                chars = page.chars  # individuele tekstfragmenten met font info
                lines = self.group_chars_into_lines(chars)
                md_lines.extend(self.convert_lines_to_md(lines))

        return "\n".join(md_lines)
    
    def group_chars_into_lines(self, chars):
        lines = {}
        for ch in chars:
            y = round(ch["top"], 1)
            lines.setdefault(y, []).append(ch)

        # sorteer op Y (boven naar beneden)
        sorted_lines = []
        for y in sorted(lines.keys()):
            line = sorted(lines[y], key=lambda c: c["x0"])
            sorted_lines.append(line)
        return sorted_lines
    
    def detect_heading_level(self, font_size):
        if font_size >= font_sizes["H1"]:
            return 1
        if font_size >= font_sizes["H2"]:
            return 2
        if font_size >= font_sizes["H3"]:
            return 3
        return None
    
    def style_text(self, text, fontname):
        if "Bold" in fontname:
            return f"**{text}**"
        if "Italic" in fontname or "Oblique" in fontname:
            return f"*{text}"
        return text
    
    def convert_lines_to_md(self, lines):
        md = []

        for line in lines:
            # combineer chars
            text = "".join(ch["text"] for ch in line).strip()
            if not text:
                md.append("")
                continue

            # detecteer heading
            avg_size = sum(ch["size"] for ch in line) / len(line)
            level = self.detect_heading_level(avg_size)
            if level:
                md.append("#" * level + " " + text)
                continue

            # detecteer bullet list
            if text.startswith(("•", "-", "◦", "‣")):
                md.append("- " + text.lstrip("•-◦‣ "))
                continue

            # detecteer bold/italic per char
            styled = "".join(self.style_text(ch["text"], ch["fontname"]) for ch in line)
            md.append(styled)

        return md
    
    def page_needs_ocr(self, page):
        text = page.extract_text()
        return not text or text.strip() == ""
    
    def ocr_page(self, page):
        # Render PDF-pagina als afbeelding
        img = page.to_image(resolution=300).original
        pil_img = Image.fromarray(img)  # gaat fou als er geen array is
        return pytesseract.image_to_string(pil_img)
    
    def pdf_to_text_with_ocr(self, path):
        pages_text = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()

                if not text or text.strip() == "":
                    # OCR fallback
                    text = self.ocr_page(page)

                pages_text.append(text or "")

        return "\n\n".join(pages_text)
    
    def pdf_to_markdown_with_ocr(self, path):
        raw_text = self.pdf_to_text_with_ocr(path)
        return self.pdf_to_markdown(raw_text)
    
    def detect_language(self, text: str) -> str:
        try:
            return detect(text)
        except LangDetectException:
            return "unknown"
        
    def ocr_page_with_lang(self, page):
        # Render PDF-pagima als afbeelding
        img = page.to_image(resolution=300).original
        pil_img = Image.fromarray(img)  # gaat fout als er geen array is

        # Probeer eerst een kleine preview voor taal detectie
        preview_text = pytesseract.image_to_string(pil_img, lang="eng")[:500]

        lang = self.detect_language(preview_text)
        tess_lang = self.LANG_MAP.get(lang, "eng")  # fallback naar Engels

        # OCR met gedetecteerde taal
        return pytesseract.image_to_string(pil_img, lang=tess_lang)

    def pdf_to_text_with_auto_ocr(self, path):
        pages_text = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()

                if not text or text.strip() == "":
                    # OCR fallback met automatische taal
                    text = self.ocr_page_with_lang(page)

                pages_text.append(text or "")

        return "\n\n".join(pages_text)

    def export_markdown_to_pdf(self, md_text, output_path):
        # 1. Markdown -> HTML
        html = render_markdown(md_text)

        # 2. HTML in QTextDocument
        doc = QTextDocument()
        doc.setHtml(html)

        # 3. PDF printer instellen
        printer = QPrinter()
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(output_path)

        # 4. Renderen
        doc.print(printer)

    def export_pdf(self):
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporteer naar PDF",
            "",
            "PDF-bestanden (*.pdf)"
        )

        if not path:
            return
        
        try:
            md_text = self.editor.toPlainText()
            self.export_markdown_to_pdf(md_text, path)
        except Exception as e:
            QMessageBox.critical(self, "Fout bij exporteren", str(e))
            return
        
        QMessageBox.information(self, "Succes", "PDF succesvol opgeslagen!")

    def html_to_docx(self, html, output_path):
        doc = Document()
        soup = BeautifulSoup(html, "html.parser")

        for el in soup.recursiveChildGenerator():
            if el.name == "h1":
                doc.add_heading(el.get_text(), level=1)
            elif el.name == "h2":
                doc.add_heading(el.get_text(), level=2)
            elif el.name == "h3":
                doc.add_heading(el.get_text(), level=3)
            elif el.name =="p":
                doc.add_paragraph(el.get_text())
            elif el.name == "pre":
                doc.add_paragraph(el.get_text(), style="Intense Quote")
            elif el.name == "code":
                doc.add_paragraph(el.get_text(), style="Intense Quote")
            elif el.name == "li":
                doc.add_paragraph(el.get_text(), style="List Bullet")
            elif el.name == "img":
                src = el["src"]
                doc.add_picture(src, width=None)  # optioneel: breedte instellen

        doc.save(output_path)

    def export_markdown_to_word(self, md_text, output_path):
        # 1. Metadata
        meta = self.extract_metadata(md_text)

        # 2. Markdown -> HTML
        html = render_markdown(md_text)

        # 3. TOC genereren
        headings = self.extract_headings(md_text)
        toc_html = self.build_clickable_toc(headings)

        # 4. TOC + body combineren
        full_html = toc_html + "<hr>" + html

        # 5. HTML -> DOCX
        self.html_to_docx(full_html, output_path)

    def extract_metadata(self, md_text):
        """
        Verwacht YAML frontmatter bovenaan het document.
        """
        fm = re.match(r"---\n(.*?)\n---", md_text, re.DOTALL)
        if not fm:
            return {}
        
        data = yaml.safe_load(fm.group(1))
        return data or {}

    def slugify(self, text):
        # Unieke maar stabiele ID
        base = re.sub(r"[^a-zA-Z0-9]+", "-", text).strip("-").lower()
        h = hashlib.md5(text.encode()).hexdigest()[:6]
        return f"{base}-{h}"

    def extract_headings(self, md_text):
        headings = []
        HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)", re.MULTILINE)
        for match in HEADING_RE.finditer(md_text):
            level = len(match.group(1))
            title = match.group(2).strip()
            anchor = self.slugify(title)
            headings.append((level, title, anchor))
        return headings

    def build_clickable_toc(self, headings):
        html = "<h1>Inhoudsopgave</h1><ul>"
        prev_level = 1

        for level, title, anchor in headings:
            if level > prev_level:
                html += "<ul>" * (level - prev_level)
            elif level < prev_level:
                html += "</ul>" * (prev_level - level)

            html += f"<li><a href=\"#{anchor}\">{title}</a></li>"
            prev_level = level

        html += "</ul>" * prev_level
        return html


    def export_word(self):
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporteer naar Word",
            "",
            "Word-bestanden (*.docx)"
        )

        if not path:
            return
        
        md_text = self.editor.toPlainText()
        self.export_markdown_to_word(md_text, path)

        QMessageBox.information(self, "Succes", "Word-document opgeslagen!")

    def export_markdown_to_epub_zonder_hoofdstukken(self, md_text, output_path):
        # 1. Metadata
        meta = self.extract_metadata(md_text)
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")

        # 2. Markdown -> HTML
        html = render_markdown(md_text)

        # 3. EPUB object aanmaken
        book = epub.EpubBook()
        book.set_identifier("id123456")
        book.set_title(title)
        book.add_author(author)
        book.set_language("nl")

        # 4. Cover afbeelding (optioneel)
        # cover_path = "path/to/cover.jpg"
        # if os.path.exists(cover_path):
        #     with open(cover_path, "rb") as f:
        #         book.set_cover("cover.jpg", f.read())

        # 5. TOC genereren
        headings = self.extract_headings(md_text)
        toc_html = self.build_clickable_toc(headings)

        toc_item = epub.EpubHtml(title="Inhoudsopgave", file_name="toc.xhtml", lang="nl")
        toc_item.content = toc_html
        book.add_item(toc_item)

        # 6. Body
        body_item = epub.EpubHtml(title=title, file_name="body.xhtml", lang="nl")
        body_item.content = html
        book.add_item(body_item)

        # TOC + spine
        book.toc = (epub.Link("toc.xhtml", "Inhoudsopgave", "toc"), epub.Link("body.xhtml", title, "body"))

        book.spine = ["nav", toc_item, body_item]

        # 8. Navigatie
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        # 9. CSS (optioneel)
        style = """
        body { font-family: Arial, sans-serif; line-height: 1.5; padding: 1em; }
        h1 { font-size: 2em; margin-top: 1em; }
        h2 { font-size: 1.5em; margin-top: 1em; }
        h3 { font-size: 1.2em; margin-top: 1em; }
        p { margin-bottom: 1em; }
        pre { font-family: "Courier New", monospace; background: #f4f4f4; padding: 1em; overflow-x: auto; }
        code { font-family: "Courier New", monospace; background: #f4f4f4; padding: 0.2em 0.4em; }
        img { max-width: 100%; height: auto; }
        """
        css = epub.EpubItem(uid="style", file_name="style.css", media_type="text/css", content=style)
        book.add_item(css)

        # 10.EPUB opslaan
        epub.write_epub(output_path, book, {})

    def export_epub(self):
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporteer naar EPUB",
            "",
            "EPUB-bestanden (*.epub)"
        )

        if not path:
            return
        
        md_text = self.editor.toPlainText()
        self.export_markdown_to_epub(md_text, path)

        QMessageBox.information(self, "Succes", "EPUB-boek opgeslagen!")

    def split_into_chapters(self, md_text):
        regels = md_text.split("\n")
        parts = []
        hoofdstuk_inhoud = ""
        for regel in regels:
            if regel.startswith("# "):
                if hoofdstuk_inhoud:
                    parts.append(hoofdstuk_inhoud)
                hoofstuktitel = regel[2:].strip()
                parts.append(hoofstuktitel)
                hoofdstuk_inhoud = ""
            else:
                hoofdstuk_inhoud += regel + "\n"
                if hoofdstuk_inhoud.strip():
                    hoofdstuk_inhoud += "\n"
        if hoofdstuk_inhoud.strip():
            parts.append(hoofdstuk_inhoud)
        

        chapters = []

        for i in range(1, len(parts), 2):
            title = parts[i].strip()
            content = parts[i+1].strip() if i+1 < len(parts) else ""
            chapters.append((title, content))

        return chapters
    
    def chapters_to_html(self, chapters):
        html_chapters = []
        for title, md in chapters:
            #html = self.markdown_to_html(md)
            html = render_markdown(md)
            html_chapters.append((title, html))
        return html_chapters

    def create_epub_chapters(self, book, html_chapters):
        epub_chapters = []

        for idx, (title, html) in enumerate(html_chapters, start=1):
            
            item = epub.EpubHtml(
                title=title,
                file_name=f"chapter_{idx}.xhtml", 
                lang="nl")
            item.content = html
            book.add_item(item)
            epub_chapters.append(item)

        return epub_chapters

    def build_epub_toc(self, epub_chapters):
        return tuple(
            epub.Link(item.file_name, item.title, f"chap{idx}")
              for idx, item in enumerate(epub_chapters, start=0)
              )
    
    def export_markdown_to_epub(self, md_text, output_path):
        meta = self.extract_metadata(md_text)
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")

        # 1. Markdown -> hoofdstukken
        chapters = self.split_into_chapters(md_text)
        html_chapters = self.chapters_to_html(chapters)

        # 2. EPUB object aanmaken
        book = epub.EpubBook()
        book.set_title(title)
        book.add_author(author)
        book.set_language("nl")

        # 3. Hoofdstukken toevoegen
        epub_items = self.create_epub_chapters(book, html_chapters)

        # 4. TOC
        book.toc = self.build_epub_toc(epub_items)

        # 5. Spine
        book.spine = ["nav"] + epub_items

        # 6. Navigatie
        book.add_item(epub.EpubNcx())

        book.add_item(epub.EpubNav())

        # remove
        book.items = [i for i in book.get_items() if i.file_name not in ('nav.xhtml','toc.ncx')]
        # add proper ones
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())


        # 7. CSS (optioneel)
        css = epub.EpubItem(uid="style", file_name="style.css", media_type="text/css", content="""
        body { font-family: Arial, sans-serif; line-height: 1.5; padding: 1em; }
        h1 { font-size: 2em; margin-top: 1em; }
        h2 { font-size: 1.5em; margin-top: 1em; }
        h3 { font-size: 1.2em; margin-top: 1em; }
        p { margin-bottom: 1em; }
        pre { font-family: "Courier New", monospace; background: #f4f4f4; padding: 1em; overflow-x: auto; }
        code { font-family: "Courier New", monospace; background: #f4f4f4; padding: 0.2em 0.4em; }
        img { max-width: 100%; height: auto; }
        """)
        book.add_item(css)

        #cover_path = self.generate_epub_cover(meta, logo_path="assets/logo.png")
        #cover_path = "assets/cover.jpg"

        #cover_path = self.generate_epub_cover_with_background(meta, bg_path="assets/background.jpg", logo_path="assets/logo.png")

        #with open(cover_path, "rb") as f:
            #book.set_cover("cover.jpg", f.read())

        # 8. EPUB opslaan
        epub.write_epub(output_path, book, {"pretty": True})            

    def closeEvent(self, event):
        if self.unsaved_changes:
            reply = QMessageBox.question(self, 'Waarschuwing', 
                                         'Huidig bestand is nog niet opgeslagen. Wil je de wijzigingen opslaan?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No | QMessageBox.StandardButton.Cancel)
            if reply == QMessageBox.StandardButton.Yes:
                self.opslaan()
                event.accept()
            elif reply == QMessageBox.StandardButton.No:
                event.accept()
            else:
                event.ignore()
        else:
            event.accept()



    def generate_epub_cover(self, meta, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560
        bg_color = (245, 245, 245)  # Lichtgrijze achtergrond

        img = Image.new("RGB", (width, height), color=bg_color)
        draw = ImageDraw.Draw(img)


        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")

        # Fonts
        #title_font = ImageFont.truetype("arial.ttf", 40)
        #author_font = ImageFont.truetype("arial.ttf", 30)
        #small_font = ImageFont.truetype("arial.ttf", 20)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titel
        #w, h = draw.textsize(title, font=title_font)
        #w, h = draw.textsize(title)
        w = draw.textlength(title)
        #draw.text(((width - w) / 2, 200), title, fill="black", font=title_font)
        draw.text(((width - w) / 2, 200), title, fill="black")

        # Auteur
        #author_font = ImageFont.truetype("arial.ttf", 30)
        #w, h = draw.textsize(author, font=author_font)
        #draw.text(((width - w) / 2, 300), author, fill="gray", font=author_font)
        #w, h = draw.textsize(author)
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 300), author, fill="gray")

        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            #w, h = draw.textsize(footer, font=small_font)
            #draw.text(((width - w) / 2, height - 50), footer, fill="gray", font=small_font)
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, height - 50), footer, fill="gray")

        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            # Schaal logo naar 20% breedte van de cover
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            # Plaats logo bovenaan
            lx = (width - logo.width) // 2
            ly = int(height * 0.10)
            img.paste(logo, (lx, ly), logo)

        img.save(output_path, "JPEG", quality=95)
        return output_path

    def zoeken_en_vervangen(self):
        #self.zoek_venster = ZoekEnVervang()
        #self.zoek_venster.show()
        ...

    def generate_epub_cover_with_background(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 200), title, fill="white")

        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 300), author, fill="lightgray")

        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, height - 50), footer, fill="lightgray")

        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = int(height * 0.10)
            bg.paste(logo, (lx, ly), logo)

        bg.save(output_path, "JPEG", quality=95)
        return output_path

    def generate_epub_cover_with_background_and_gradient(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 200), title, fill="white")

        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 300), author, fill="lightgray")

        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, height - 50), footer, fill="lightgray")

        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = int(height * 0.10)
            bg.paste(logo, (lx, ly), logo)

        bg.save(output_path, "JPEG", quality=95)
        return output_path
    
    def generate_epub_cover_with_background_and_gradient_and_title_block(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titelblok achtergrond
        block_height = 400
        block_color = (30, 30, 30, 200)  # Semi-transparant donkergrijs
        block = Image.new("RGBA", (width, block_height), block_color)
        bg.paste(block, (0, 200), block)
        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 250), title, fill="white")
        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 350), author, fill="lightgray")
        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, 550), footer, fill="lightgray")
        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = 200 + (block_height - logo.height) // 2
            bg.paste(logo, (lx, ly), logo)
        bg.save(output_path, "JPEG", quality=95)
        return output_path
    
    def generate_epub_cover_with_background_and_gradient_and_title_block_and_subtitle(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        subtitle = meta.get("subtitle", "")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titelblok achtergrond
        block_height = 500
        block_color = (30, 30, 30, 200)  # Semi-transparant donkergrijs
        block = Image.new("RGBA", (width, block_height), block_color)
        bg.paste(block, (0, 200), block)
        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 250), title, fill="white")
        # Subtitel
        if subtitle:
            w = draw.textlength(subtitle)
            draw.text(((width - w) / 2, 320), subtitle, fill="lightgray")
        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 400), author, fill="lightgray")
        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, 550), footer, fill="lightgray")
        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = 200 + (block_height - logo.height) // 2
            bg.paste(logo, (lx, ly), logo)
        bg.save(output_path, "JPEG", quality=95)
        return output_path
    
    def generate_epub_cover_with_background_and_gradient_and_title_block_and_subtitle_and_author(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        subtitle = meta.get("subtitle", "")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titelblok achtergrond
        block_height = 500
        block_color = (30, 30, 30, 200)  # Semi-transparant donkergrijs
        block = Image.new("RGBA", (width, block_height), block_color)
        bg.paste(block, (0, 200), block)
        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 250), title, fill="white")
        # Subtitel
        if subtitle:
            w = draw.textlength(subtitle)
            draw.text(((width - w) / 2, 320), subtitle, fill="lightgray")
        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 400), author, fill="lightgray")
        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, 550), footer, fill="lightgray")
        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = 200 + (block_height - logo.height) // 2
            bg.paste(logo, (lx, ly), logo)
        bg.save(output_path, "JPEG", quality=95)
        return output_path

    def generate_epub_cover_with_background_and_gradient_and_title_block_and_subtitle_and_author_and_project(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        subtitle = meta.get("subtitle", "")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titelblok achtergrond
        block_height = 500
        block_color = (30, 30, 30, 200)  # Semi-transparant donkergrijs
        block = Image.new("RGBA", (width, block_height), block_color)
        bg.paste(block, (0, 200), block)
        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 250), title, fill="white")
        # Subtitel
        if subtitle:
            w = draw.textlength(subtitle)
            draw.text(((width - w) / 2, 320), subtitle, fill="lightgray")
        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 400), author, fill="lightgray")
        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, 550), footer, fill="lightgray")
        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = 200 + (block_height - logo.height) // 2
            bg.paste(logo, (lx, ly), logo)
        bg.save(output_path, "JPEG", quality=95)
        return output_path

    def generate_epub_cover_with_background_and_gradient_and_title_block_and_subtitle_and_author_and_project_and_version(self, meta, bg_path, logo_path=None, output_path="cover.jpg"):
        # Afmetingen volgens EPUB-conventies
        width, height = 1600, 2560

        # Achtergrond
        if os.path.exists(bg_path):
            bg = Image.open(bg_path).convert("RGB")
            bg = bg.resize((width, height), Image.LANCZOS)
        else:
            bg = Image.new("RGB", (width, height), color=(245, 245, 245))

        # Gradient overlay
        gradient = Image.new("L", (1, height), color=0xFF)
        for y in range(height):
            gradient.putpixel((0, y), int(255 * (1 - y / height)))  # van wit naar transparant
        alpha_gradient = gradient.resize((width, height))
        black_img = Image.new("RGBA", (width, height), color=(0, 0, 0, 255))
        black_img.putalpha(alpha_gradient)
        bg = Image.alpha_composite(bg.convert("RGBA"), black_img)

        draw = ImageDraw.Draw(bg)

        # Metadata
        title = meta.get("title", "Mijn Markdown Boek")
        subtitle = meta.get("subtitle", "")
        author = meta.get("author", "Onbekende Auteur")
        project = meta.get("project", "")
        version = meta.get("version", "")

        # Titelblok achtergrond
        block_height = 500
        block_color = (30, 30, 30, 200)  # Semi-transparant donkergrijs
        block = Image.new("RGBA", (width, block_height), block_color)
        bg.paste(block, (0, 200), block)
        # Titel
        w = draw.textlength(title)
        draw.text(((width - w) / 2, 250), title, fill="white")
        # Subtitel
        if subtitle:
            w = draw.textlength(subtitle)
            draw.text(((width - w) / 2, 320), subtitle, fill="lightgray")
        # Auteur
        w = draw.textlength(author)
        draw.text(((width - w) / 2, 400), author, fill="lightgray")
        # Project + versie
        footer = f"{project} {version}".strip()
        if footer:
            w = draw.textlength(footer)
            draw.text(((width - w) / 2, 550), footer, fill="lightgray")
        # Logo (optioneel)
        if logo_path and os.path.exists(logo_path):
            logo = Image.open(logo_path).convert("RGBA")
            target_w = int(width * 0.2)
            aspect = logo.height / logo.width
            logo = logo.resize((target_w, int(target_w * aspect)), Image.LANCZOS)

            lx = (width - logo.width) // 2
            ly = 200 + (block_height - logo.height) // 2
            bg.paste(logo, (lx, ly), logo)
        bg.save(output_path, "JPEG", quality=95)
        return output_path
    
    def extract_epub_html(self, epub_path):
        book = epub.read_epub(epub_path)
        html_items = []

        for item in book.get_items():
            #print(f"Item: {item.get_name()} - {item.get_type()}")
            ct = self.content_type(item)
            #print(f"Content type: {ct}")

            if ct == "application/xhtml+xml":
                html_items.append((item.get_name(), item.get_content().decode("utf-8")))

        return html_items, book

    def content_type(self, item): 
        if hasattr(item, 'get_content_type'): 
            return item.get_content_type() 
        if hasattr(item, 'media_type'): 
            return item.media_type 
        t = getattr(item, 'get_type', None) 
        return t() if callable(t) else None
    
    def epub_metadata_to_frontmatter(self, book):
        title = book.get_metadata("DC", "title")[0][0] if book.get_metadata("DC", "title") else "Mijn Markdown Boek"
        author = book.get_metadata("DC", "creator")[0][0] if book.get_metadata("DC", "creator") else "Onbekende Auteur"
        language = book.get_metadata("DC", "language")[0][0] if book.get_metadata("DC", "language") else "nl"
        identifier = book.get_metadata("DC", "identifier")[0][0] if book.get_metadata("DC", "identifier") else "id123456"

        fm = f"""---
title: {title}
author: {author}
language: {language}
identifier: {identifier}
---
"""
        return fm
    
    def html_to_markdown(self, html):
        return md(html, heading_style="ATX")  #  heading zonder de "s"
    
    def extract_images(self, book, output_dir):
        mapping = {}

        for item in book.get_items():
            ct = self.content_type(item)
            #print(f"Item: {item.get_name()} - {item.get_type()} - {ct}")
            if ct == "image/jpeg" or ct == "image/png" or ct.startswith("image/"):
                filename = os.path.basename(item.file_name)
                output_path = os.path.join(output_dir, filename)
                with open(output_path, "wb") as f:
                    f.write(item.get_content())
                mapping[item.file_name] = output_path
        return mapping

    def rewrite_image_paths(self, md_text, mapping):
        for epub_path, local_path in mapping.items():
            md_text = md_text.replace(epub_path, local_path)
        return md_text

    def epub_to_markdown(self, epub_path, output_dir="imported_epub"):
        html_items, book = self.extract_epub_html(epub_path)
        fm = self.epub_metadata_to_frontmatter(book)

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # afbeeldingen
        img_map = self.extract_images(book, os.path.join(output_dir, "images"))

        # hoofdstukken converteren
        chapters_md = []
        for filename, html in html_items:
            # md_text = self.html_to_markdown(html)
            md_text = self.html_with_footnotes_to_markdown(html)
            md_text = self.rewrite_image_paths(md_text, img_map)
            chapters_md.append(md_text)

        # samenvoegen
        full_md = fm + "\n\n" + "\n\n".join(chapters_md)
        return full_md

    def import_epub(self):
        path, _ = QFileDialog.getOpenFileName(
            self,
            "Importeer EPUB",
            "",
            "EPUB-bestanden (*.epub)"
        )

        if not path:
            return
        
        md_text = self.epub_to_markdown(path)
        self.editor.setPlainText(md_text)
        QMessageBox.information(self, "Succes", "EPUB-boek geïmporteerd!")

    def extract_footnotes_from_html(self, html):
        soup = BeautifulSoup(html, "html.parser")
        notes = {}

        # typische structuur: <a href="#fn1">1</a> ... <div id="fn1">Voetnoot tekst</div>
        for ref in soup.find_all("a", href=True):
            href = ref["href"]
            if ref["href"].startswith("#"):
                target_id = href[1:]
                target = soup.find(id=target_id)
                if target:
                    num = ref.get_text(strip=True)
                    text = target.get_text(" ", strip=True)
                    notes[num] = text
        return notes
    
    def replace_refs_with_markdown_footnotes(self, html):
        soup = BeautifulSoup(html, "html.parser")

        footnotes = self.extract_footnotes_from_html(html)

        for ref in soup.find_all("a", href=True):
            href = ref["href"]
            if href.startswith("#"):
                num = ref.get_text(strip=True)
                if num in footnotes:
                    md_ref = f"[^{num}]"
                    ref.replace_with(md_ref)

        return str(soup)

    def html_with_footnotes_to_markdown(self, html):
        # 1. footnotes extraheren en refs vervangen door markdown syntax
        notes = self.extract_footnotes_from_html(html)

        # 2. refs in HTML vervangen door markdown syntax
        html_clean = self.replace_refs_with_markdown_footnotes(html)

        # 3. HTML -> Markdown
        body_md = self.html_to_markdown(html_clean)

        # 4. footnotes toevoegen aan einde van markdown
        if notes:
            body_md += "\n\n"
            for num, text in notes.items():
                body_md += f"[^{num}]: {text}\n"

        return body_md

    # zoeken en vervangen

    def find_next(self):
        text = self.find_input.text()    
        if not text:
            return
        doc = self.editor.document()    
        cursor = self.editor.textCursor()    
        start_pos = cursor.selectionEnd() if cursor.hasSelection() else cursor.position()
        if self.case_cb.isChecked():
            # plain search from int position
            it = doc.find(text, start_pos)        
            if it.isNull():            
                it = doc.find(text, 0)    
            else:        
                # regex with case-insensitive option        
                regex = QRegularExpression(text)        
                regex.setPatternOptions(QRegularExpression.PatternOption.CaseInsensitiveOption)
                it = doc.find(regex, start_pos)        
                if it.isNull():            
                    it = doc.find(regex, 0)
            if not it.isNull():        
                self.editor.setTextCursor(it)        
                self.update_highlight()


    def find_previous(self):
        # backwards search: iterate matches and pick last before current position
        text = self.find_input.text()
        if not text:
            return
        doc = self.editor.document()

        cur = self.editor.textCursor()
        pos = cur.selectionStart() if cur.hasSelection() else cur.position()

        if self.case_cb.isChecked():
            it = doc.find(text, 0)
        else:
            regex = QRegularExpression(text)
            regex.setPatternOptions(QRegularExpression.PatternOption.CaseInsensitiveOption)
            it = doc.find(regex, 0)

        last = None
        while not it.isNull() and it.selectionEnd() <= pos:
            last = it
            it = doc.find(text, it.selectionEnd())
        if last:
            self.editor.setTextCursor(last)
        else:
            # wrap to last match in document
            it = doc.find(text, 0)
            last = None
            while not it.isNull():
                last = it
                it = doc.find(text, it.selectionEnd())
            if last:
                self.editor.setTextCursor(last)
        self.update_highlight()

    def replace_one(self):
        cursor = self.editor.textCursor()
        find_text = self.find_input.text()
        if not find_text:
            return
        selected = cursor.selectedText()
        if cursor.hasSelection() and ((self.case_cb.isChecked() and selected == find_text) or (not self.case_cb.isChecked() and selected.lower() == find_text.lower())):
            cursor.insertText(self.replace_input.text())
            self.editor.setTextCursor(cursor)
        self.find_next()

    def replace_all(self):
        find_text = self.find_input.text()
        if not find_text:
            return
        text = self.editor.toPlainText()
        if self.case_cb.isChecked():
            new_text = text.replace(find_text, self.replace_input.text())
        else:
            # case-insensitive replace preserving original case is complex; use simple lower-replace
            import re
            pattern = re.compile(re.escape(find_text), re.IGNORECASE)
            new_text = pattern.sub(self.replace_input.text(), text)
        self.editor.setPlainText(new_text)
        self.update_highlight()

    def update_highlight(self):
        # Clear selections and add new extra selections for all matches.
        find_text = self.find_input.text()
        extra_selections = []

        if find_text:
            doc = self.editor.document()
            #flags = 0 if self.case_cb.isChecked() else getattr(__import__('PyQt6.Qt', fromlist=['Qt']).Qt, 'CaseInsensitive', 0)
            flags = 0
            it = doc.find(find_text, 0)
            cursors = []
            while not it.isNull():
                cursors.append(it)
                it = doc.find(find_text, it.selectionEnd())

            # Create ExtraSelection objects
            for c in cursors:
                sel = QTextCursor(c)
                selection = QTextCursor(sel)
                extra = self._make_extra(selection, self.highlight_format)
                extra_selections.append(extra)

            # If current cursor is on a match, make it highlighted differently
            cur = self.editor.textCursor()
            for i, c in enumerate(cursors):
                if c.selectionStart() <= cur.position() <= c.selectionEnd():
                    # mark this one with match_format
                    extra_selections[i] = self._make_extra(QTextCursor(c), self.match_format)
                    break

        self.editor.setExtraSelections(extra_selections)

    def _make_extra(self, cursor: QTextCursor, fmt: QTextCharFormat):
        from PyQt6.QtWidgets import QTextEdit
        extra = QTextEdit.ExtraSelection()
        extra.cursor = cursor
        extra.format = fmt
        return extra

    def markdown_to_plain_text(self, md_text, strip_frontmatter=True):

        if strip_frontmatter:
            md_text = re.sub(r"^---\n.*?\n---\n", "", md_text, flags=re.DOTALL)
        html = markdown(md_text)
        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text("\n")

        # Opschonen van meerdere nieuwe regels
        lines = [line.strip() for line in text.splitlines()]

        return "\n".join(lines).strip()
    
    def export_txt(self):
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Exporteer als Tekstbestand",
            "",
            "Tekstbestanden (*.txt)"
        )

        if not path:
            return
        
        md_text = self.editor.toPlainText()
        plain_text = self.markdown_to_plain_text(md_text)

        try:
            with open(path, "w", encoding="utf-8") as f:
                f.write(plain_text)
        except Exception as e:
            QMessageBox.critical(self, "Fout", f"Fout bij exporteren: {e}")
        else:
            QMessageBox.information(self, "Succes", "Bestand geëxporteerd als tekstbestand!")

    def spellcheck(self):
        QMessageBox.information(self, "Spellcheck", "Spelling controleren...\nDit kan lang duren, je kunt meldingen krijgen dat het programma niet meer reageert. \nHeb dan geduld.\nAls het klaar is krijg je daar een melding van.")
        if not hasattr(self, "spellchecker"):
            self.spellchecker = SpellChecker()
        text = self.editor.toPlainText()
        matches = self.spellchecker.check(text)
        verslag = "Spellingcontrole verslag:\n\n"
        for match in matches:
            line, column = match.get_line_and_column(text)
            cursor = self.editor.textCursor()
            block = self.editor.document().findBlockByNumber(line - 1)
            if block.isValid():
                pos = block.position() + column - 1
            start = pos
            end = pos + match.error_length
            word = self.get_word_from_line_column(line, column, match.error_length)
            verslag += f"Fout: regel:{line} kolom:{column} {word}\nSuggesties: {match.replacements[:5]}\n"

            cursor.setPosition(start)
            cursor.setPosition(end, QTextCursor.MoveMode.KeepAnchor)
            fmt = cursor.charFormat()
            fmt.setUnderlineColor(Qt.GlobalColor.red)
            fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.SpellCheckUnderline)
            cursor.setCharFormat(fmt)
        
        verslaglocatie = configuratie["opslaglocatie"] + "/spellcheck_report.txt"
        with open(verslaglocatie, "w") as f:
            f.write(verslag)
        QMessageBox.information(self, "Spellcheck", f"Spellingcontrole voltooid! {len(matches)} fouten gevonden.\nVerslag opgeslagen in {verslaglocatie}")

    def go_to_line_column(self, line, column):
        cursor = self.editor.textCursor()
        block = self.editor.document().findBlockByNumber(line - 1)
        if block.isValid():
            pos = block.position() + column - 1
            cursor.setPosition(pos)
            self.editor.setTextCursor(cursor)
            self.editor.setFocus()

    def get_word_from_line_column(self, line, column, length):
        cursor = self.editor.textCursor()
        block = self.editor.document().findBlockByNumber(line - 1)
        if block.isValid():
            pos = block.position() + column - 1
            cursor.setPosition(pos)
            cursor.setPosition(pos + length, QTextCursor.MoveMode.KeepAnchor)
            return cursor.selectedText()
        return ""


class SpellChecker:
    def __init__(self, lang="nl"):
        self.tool = language_tool_python.LanguageTool(lang)
        

    def check(self,text):
        return self.tool.check(text)
    
    def suggestions_for(self, word):
        matches = self.tool.check(word)
        if matches:
            return matches[0].replacements
        return []    